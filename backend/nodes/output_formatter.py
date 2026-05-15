"""Output formatter node.

The pipeline stores markdown. PDF and Word conversion stay available for the
download endpoint, where users choose the desired format after generation.
"""

from io import BytesIO
from typing import Any, Dict, List

from backend.classes.state import ResearchState


async def output_formatter_node(state: ResearchState) -> Dict[str, Any]:
    report = state.get("report", "")
    events: List[Dict[str, Any]] = [
        {"type": "status", "node": "output_formatter", "message": "Formatting markdown output"},
        {"type": "status", "node": "output_formatter", "message": "Markdown output ready"},
    ]
    return {"output": report, "events": events}


def _to_pdf(markdown_text: str, title: str) -> bytes:
    from backend.services.pdf_service import PDFService

    ok, result = PDFService().generate_pdf_bytes(markdown_text, title, language="zh")
    if ok:
        return result
    raise RuntimeError(f"PDF generation failed: {result}")


def _set_cjk_font(style, font_name: str, pt_size=None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    style.font.name = font_name
    rpr = style.element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), font_name)
    if pt_size:
        style.font.size = Pt(pt_size)


def _to_docx(markdown_text: str) -> bytes:
    from docx import Document

    doc = Document()

    _set_cjk_font(doc.styles["Normal"], "Source Han Sans SC", pt_size=10.5)
    for level, pt in ((1, 18), (2, 14), (3, 12)):
        heading_style = doc.styles[f"Heading {level}"]
        _set_cjk_font(heading_style, "Source Han Sans SC", pt_size=pt)

    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("|") and stripped.endswith("|"):
            doc.add_paragraph(stripped)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
